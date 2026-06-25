from fpdf import FPDF
from datetime import datetime
from typing import List, Any

class ChatPDF(FPDF):
    def header(self):
        # Header title
        self.set_font('helvetica', 'B', 14)
        self.cell(0, 10, 'AI Chatbot Assistant - Chat History', border=False, align='C')
        self.ln(12)
        
    def footer(self):
        # Position at 1.5 cm from bottom
        self.set_y(-15)
        self.set_font('helvetica', 'I', 8)
        # Page number
        self.cell(0, 10, f'Page {self.page_no()}', border=False, align='C')

class ExportHelpers:
    @staticmethod
    def export_to_text(chat_title: str, messages: List[Any]) -> str:
        timestamp = datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S UTC")
        lines = [
            "==================================================",
            f"CHAT HISTORY EXPORT",
            f"Session Title: {chat_title}",
            f"Exported At: {timestamp}",
            "==================================================",
            "\n"
        ]

        for msg in messages:
            sender_label = "USER" if msg.sender == "user" else "AI ASSISTANT"
            msg_time = msg.timestamp.strftime("%Y-%m-%d %H:%M:%S")
            lines.append(f"[{sender_label}] ({msg_time}):")
            lines.append(msg.content)
            lines.append("-" * 50)
            lines.append("")

        return "\n".join(lines)

    @staticmethod
    def export_to_pdf(chat_title: str, messages: List[Any]) -> bytes:
        pdf = ChatPDF()
        pdf.add_page()
        pdf.set_font("helvetica", "B", 12)
        
        # Topic
        pdf.cell(0, 8, f"Topic: {chat_title}", border="B", ln=True)
        pdf.ln(5)
        
        for msg in messages:
            sender_label = "User" if msg.sender == "user" else "AI Assistant"
            msg_time = msg.timestamp.strftime("%Y-%m-%d %H:%M")
            
            # Print Speaker Header
            pdf.set_font("helvetica", "B", 10)
            if msg.sender == "user":
                pdf.set_text_color(26, 115, 232)  # Blue for user
            else:
                pdf.set_text_color(16, 124, 65)   # Green for AI
                
            pdf.cell(0, 6, f"{sender_label} ({msg_time})", ln=True)
            
            # Print Message Content
            pdf.set_text_color(0, 0, 0)
            pdf.set_font("helvetica", "", 10)
            
            # Replace characters that are outside latin-1 to avoid fpdf errors
            content_cleaned = msg.content.encode("latin-1", "replace").decode("latin-1")
            
            # Use multi_cell to handle wrapping text
            pdf.multi_cell(0, 5, content_cleaned)
            pdf.ln(4)
            
        return bytes(pdf.output())

    @staticmethod
    def export_to_docx(chat_title: str, messages: List[Any]) -> bytes:
        import io
        from docx import Document as DocxDocument
        from docx.shared import Pt, RGBColor
        
        doc = DocxDocument()
        
        # Title Header
        title_p = doc.add_paragraph()
        title_run = title_p.add_run(f"AI Chatbot Assistant - Chat History\nTopic: {chat_title}")
        title_run.font.name = 'Arial'
        title_run.font.size = Pt(14)
        title_run.bold = True
        
        # Subtitle
        timestamp = datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S UTC")
        subtitle_p = doc.add_paragraph()
        subtitle_run = subtitle_p.add_run(f"Exported At: {timestamp}\n")
        subtitle_run.font.name = 'Arial'
        subtitle_run.font.size = Pt(10)
        subtitle_run.italic = True
        subtitle_run.font.color.rgb = RGBColor(128, 128, 128)
        
        doc.add_paragraph("=" * 60)
        
        for msg in messages:
            sender_label = "User" if msg.sender == "user" else "AI Assistant"
            msg_time = msg.timestamp.strftime("%Y-%m-%d %H:%M")
            
            # Speaker paragraph
            speaker_p = doc.add_paragraph()
            speaker_run = speaker_p.add_run(f"{sender_label} ({msg_time})")
            speaker_run.font.name = 'Arial'
            speaker_run.font.size = Pt(11)
            speaker_run.bold = True
            
            if msg.sender == "user":
                speaker_run.font.color.rgb = RGBColor(26, 115, 232)
            else:
                speaker_run.font.color.rgb = RGBColor(16, 124, 65)
                
            # Content paragraph
            content_p = doc.add_paragraph()
            content_run = content_p.add_run(msg.content)
            content_run.font.name = 'Arial'
            content_run.font.size = Pt(10.5)
            
            doc.add_paragraph("-" * 50)
            
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream.getvalue()
